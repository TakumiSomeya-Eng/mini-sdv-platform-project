#!/usr/bin/env python3
"""
Interview Document Generator (python-docx version)
Usage: python generate_interview_py.py <FR_NUMBER> <FEATURE_TITLE> <PROJECT_NAME> [OUTPUT_DIR]
"""

import sys
import os
import re
from datetime import date
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    return p


def add_placeholder(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(f"[ {text} ]")
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run.font.italic = True
    return p


def add_table_row(table, cells):
    row = table.add_row()
    for i, text in enumerate(cells):
        row.cells[i].text = text
        for para in row.cells[i].paragraphs:
            para.runs[0].font.size = Pt(10) if para.runs else None


def generate(fr_number, fr_title, project_name, output_dir):
    doc = Document()

    # ── Page margins ──────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.2)
        section.right_margin  = Inches(1.2)

    # ── Title ─────────────────────────────────────────────────────────────────
    title = doc.add_heading(f"{fr_number}: {fr_title}", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(f"Date: {date.today()}  |  Project: {project_name}")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()

    # ── 1. Background & Challenge ─────────────────────────────────────────────
    set_heading(doc, "1. Background & Challenge")

    set_heading(doc, "Project Context", level=2)
    add_placeholder(doc, "Insert project background, business context, and problem statement here")

    set_heading(doc, "Problem Statement", level=2)
    add_placeholder(doc, "Specific challenge to be solved")
    add_placeholder(doc, "Expected outcome or success criteria")

    # ── 2. Solution & Approach ────────────────────────────────────────────────
    set_heading(doc, "2. Solution & Approach")

    set_heading(doc, "Proposed Solution", level=2)
    add_placeholder(doc, "High-level overview of the solution")

    set_heading(doc, "Implementation Strategy", level=2)
    add_placeholder(doc, "Key architectural decision #1")
    add_placeholder(doc, "Key architectural decision #2")

    # ── 3. Implementation Details ─────────────────────────────────────────────
    set_heading(doc, "3. Implementation Details")

    set_heading(doc, "Tech Stack", level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for cell, text in zip(hdr, ["Technology", "Version", "Rationale"]):
        cell.text = text
        for para in cell.paragraphs:
            para.runs[0].bold = True
            para.runs[0].font.size = Pt(10)

    for tech in [("[ Technology ]", "[ version ]", "[ reason ]"),
                 ("[ Technology ]", "[ version ]", "[ reason ]"),
                 ("[ Technology ]", "[ version ]", "[ reason ]")]:
        add_table_row(table, tech)

    doc.add_paragraph()
    set_heading(doc, "Key Features", level=2)
    for feat in ["Feature 1 — brief description", "Feature 2 — brief description"]:
        add_placeholder(doc, feat)

    # ── 4. Demonstration & Usage ──────────────────────────────────────────────
    set_heading(doc, "4. Demonstration & Usage")

    set_heading(doc, "Setup Instructions", level=2)
    for step in [
        "Clone the repository and follow the README instructions",
        "Install dependencies",
        "Configure environment variables (see README)",
    ]:
        p = doc.add_paragraph(style="List Number")
        p.add_run(step)

    set_heading(doc, "Demo Workflow", level=2)
    add_placeholder(doc, "Step 1 of demo")
    add_placeholder(doc, "Step 2 of demo")
    add_placeholder(doc, "Expected output / success criteria")

    # ── 5. FAQ & Technical Questions ──────────────────────────────────────────
    set_heading(doc, "5. FAQ & Technical Questions")

    for i in range(1, 4):
        q = doc.add_paragraph()
        q.add_run(f"Q{i}: ").bold = True
        q.add_run("[ Question ]")
        a = doc.add_paragraph()
        a.add_run("A:  ").bold = True
        run = a.add_run("[ Answer ]")
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        run.font.italic = True
        doc.add_paragraph()

    # ── 6. Future Improvements ────────────────────────────────────────────────
    set_heading(doc, "6. Future Improvements")
    add_placeholder(doc, "Improvement 1")
    add_placeholder(doc, "Improvement 2")

    # ── Appendix ──────────────────────────────────────────────────────────────
    set_heading(doc, "Appendix: Code References")

    set_heading(doc, "Repository Structure", level=2)
    add_placeholder(doc, "Key directories and files")

    set_heading(doc, "Related Documentation", level=2)
    add_placeholder(doc, "Links to PRD / FRD / TRD / architecture review")

    # ── Save ──────────────────────────────────────────────────────────────────
    os.makedirs(output_dir, exist_ok=True)
    filename = f"{fr_number}_interview_{re.sub(r'[^a-z0-9]+', '_', fr_title.lower()).strip('_')}.docx"
    output_path = os.path.join(output_dir, filename)
    doc.save(output_path)
    print(f"[OK] {output_path}")
    return output_path


if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("Usage: python generate_interview_py.py <FR_NUMBER> <FEATURE_TITLE> [PROJECT_NAME] [OUTPUT_DIR]")
        sys.exit(1)

    fr_number    = sys.argv[1]
    fr_title     = sys.argv[2]
    project_name = sys.argv[3] if len(sys.argv) > 3 else "mini-sdv-platform"
    output_dir   = sys.argv[4] if len(sys.argv) > 4 else os.path.join(os.path.dirname(__file__), "..", "interview")

    generate(fr_number, fr_title, project_name, output_dir)
